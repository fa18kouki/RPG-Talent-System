#!/usr/bin/env python3
"""Quest HR - Research Reports Generator"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)
    return heading


def add_table_row(table, cells_text, is_header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.style.font.size = Pt(10)
        if is_header:
            set_cell_shading(cell, '2D1B69')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
    return row


def create_competitive_analysis():
    """競合調査レポート"""
    doc = Document()

    # Title
    title = doc.add_heading('Quest HR 競合調査レポート', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)

    doc.add_paragraph('作成日: 2026年3月1日')
    doc.add_paragraph('対象: Quest HR（RPG型ゲーミフィケーション人材育成プラットフォーム）')
    doc.add_paragraph('')

    # ─── 1. エグゼクティブサマリー ───
    add_styled_heading(doc, '1. エグゼクティブサマリー')
    doc.add_paragraph(
        'Quest HRは、RPGゲーミフィケーションを活用した人材育成・タレントマネジメントプラットフォームである。'
        '従業員をRPGキャラクターとして表現し、業務課題を「クエスト」として提示することで、'
        'スキル開発とエンゲージメント向上を同時に実現する。'
    )
    doc.add_paragraph(
        '本レポートでは、グローバルおよび国内のゲーミフィケーション×HR市場における主要競合を分析し、'
        'Quest HRのポジショニングと差別化戦略を明確にする。'
    )

    # ─── 2. 市場概況 ───
    add_styled_heading(doc, '2. 市場概況')

    add_styled_heading(doc, '2.1 グローバル市場', level=2)
    doc.add_paragraph(
        'ゲーミフィケーション市場全体は2025年時点で約170億～291億ドル規模と推定され、'
        '2031年までに1,000億ドルを超える成長が見込まれている（CAGR 25～28%）。'
        'うちHR・研修向けゲーミフィケーションソフトウェア市場は2025年時点で約12億ドル規模であり、'
        '2033～2034年までに約24億ドルに倍増する予測（CAGR 8.3%）。'
    )
    bullets = [
        'Fortune 500企業の78%以上がゲーミフィケーションソリューションを導入済み',
        'HR・研修分野は最も成長が速いアプリケーションセグメント（CAGR 27.9%）',
        'クラウドベースのプラットフォームが2025年収益の67.62%を占有',
        'アジア太平洋地域が最も高い成長率を示す（CAGR 28.6%）',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    add_styled_heading(doc, '2.2 国内市場', level=2)
    doc.add_paragraph(
        '日本国内のゲーミフィケーション市場は2024年時点で463億円と推定されている'
        '（セガ エックスディー調査）。業界カオスマップには216社の事業者が掲載されており、'
        '2023年度の198社から18社増加している。特に「ヘルスケア」「業務支援/ゲーミフィケーションツール」'
        'カテゴリーでの増加が顕著である。'
    )
    doc.add_paragraph(
        '経済産業省もゲーミフィケーションを活用した人材育成の有効性に着目し、'
        'デジタル人材育成におけるゲーミフィケーション活用の調査事業を実施している。'
    )

    # ─── 3. 主要競合分析 ───
    add_styled_heading(doc, '3. 主要競合分析（グローバル）')

    # Table
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    headers = ['企業/製品名', '主な特徴', '対象市場', '価格帯', '強み/弱み']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], '2D1B69')
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    competitors = [
        ['Centrical', 'リアルタイムパフォーマンス追跡、マイクロラーニング、AI駆動コーチング、ゲーミフィケーション',
         'エンタープライズ', '要問い合わせ（年間契約）',
         '強み: 包括的な機能、大企業実績\n弱み: 高コスト、導入期間長い'],
        ['Rallyware', '個別化された日次アクションプラン、業務連動型ゲーミフィケーション',
         'エンタープライズ／中堅', '要問い合わせ',
         '強み: 業績連動型、自動化エンジン\n弱み: セットアップ複雑'],
        ['Spinify', 'セールス特化型ゲーミフィケーション、リアルタイムリーダーボード、CRM連携',
         'セールスチーム', '$15～/ユーザー/月',
         '強み: CRM連携充実\n弱み: セールス特化で汎用性低い'],
        ['Engagedly', 'AIエージェント、LXP、ゲーミフィケーション、コンプライアンス追跡',
         '中堅～大企業', '$6～/ユーザー/月',
         '強み: AI統合、学習パス\n弱み: UI/UXが複雑'],
        ['Bravon', 'オールインワンゲーミフィケーションプラットフォーム、エンゲージメント向上',
         'SMB～中堅', '要問い合わせ',
         '強み: 幅広い適用範囲\n弱み: ブランド認知度低い'],
        ['The Talent Games', 'AI駆動型ゲーム化アセスメント、採用特化',
         'エンタープライズ', '要問い合わせ',
         '強み: 採用プロセスに強い\n弱み: 人材育成機能が限定的'],
        ['HiBob', '総合HR管理、Kudos/Shoutouts、エンゲージメント調査',
         '中堅～大企業', '要問い合わせ',
         '強み: HRプラットフォームとしての総合力\n弱み: ゲーミフィケーションは補助的'],
    ]

    for comp in competitors:
        row = table.add_row()
        for i, val in enumerate(comp):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph('')

    # ─── 3.2 国内競合 ───
    add_styled_heading(doc, '3.2 国内主要プレイヤー', level=2)

    table2 = doc.add_table(rows=1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'

    hdr2 = table2.rows[0].cells
    headers2 = ['企業名', '事業内容', 'カテゴリー', '特記事項']
    for i, h in enumerate(headers2):
        hdr2[i].text = h
        set_cell_shading(hdr2[i], '2D1B69')
        for p in hdr2[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    jp_competitors = [
        ['セガ エックスディー', 'ゲーミフィケーション受託・SaaS', '総合ゲーミフィケーション',
         '業界カオスマップ制作、市場調査リーダー。2025年にゲーミフィケーション研究所設立'],
        ['日本ゲーミフィケーション協会', '啓蒙・認定・コンサルティング', '業界団体',
         '認定資格発行、セミナー開催'],
        ['体験型研修企業群', 'チームビルディング、リーダーシップ研修', '体験型研修',
         'オフライン中心、デジタル移行が課題'],
        ['eラーニング×ゲーミフィケーション', 'LMS内にゲーム要素を組み込み', '専門知識研修',
         'バッジ・ポイント等の基本要素中心'],
    ]

    for comp in jp_competitors:
        row = table2.add_row()
        for i, val in enumerate(comp):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ─── 4. 競合比較マトリクス ───
    add_styled_heading(doc, '4. Quest HR 競合比較マトリクス')

    table3 = doc.add_table(rows=1, cols=7)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Table Grid'

    hdr3 = table3.rows[0].cells
    headers3 = ['機能', 'Quest HR', 'Centrical', 'Rallyware', 'Engagedly', 'Spinify', 'HiBob']
    for i, h in enumerate(headers3):
        hdr3[i].text = h
        set_cell_shading(hdr3[i], '2D1B69')
        for p in hdr3[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)

    features = [
        ['RPG型キャラクター育成',      'O', 'X', 'X', 'X', 'X', 'X'],
        ['クエスト型タスク管理',        'O', 'Triangle', 'O', 'Triangle', 'Triangle', 'X'],
        ['スキルレーダーチャート',      'O', 'O', 'Triangle', 'O', 'X', 'Triangle'],
        ['AIチャット（日次振り返り）',  'O', 'O', 'X', 'O', 'X', 'X'],
        ['アバターカスタマイズ',        'O', 'X', 'X', 'X', 'X', 'X'],
        ['レベル/XPシステム',           'O', 'O', 'O', 'O', 'O', 'X'],
        ['ファイルアップロード提出',    'O', 'O', 'O', 'O', 'X', 'X'],
        ['多言語対応（日本語）',        'O', 'X', 'X', 'X', 'X', 'X'],
        ['リアルタイム分析',            'Triangle', 'O', 'O', 'O', 'O', 'O'],
        ['CRM/外部システム連携',        'X', 'O', 'O', 'O', 'O', 'O'],
        ['モバイルアプリ',              'X', 'O', 'O', 'O', 'O', 'O'],
    ]

    symbols = {'O': '◎', 'Triangle': '△', 'X': '✕'}
    for feat in features:
        row = table3.add_row()
        for i, val in enumerate(feat):
            display = symbols.get(val, val)
            row.cells[i].text = display
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph('◎ = 対応  △ = 部分対応  ✕ = 未対応')

    # ─── 5. SWOT分析 ───
    add_styled_heading(doc, '5. Quest HR SWOT分析')

    add_styled_heading(doc, '強み (Strengths)', level=2)
    strengths = [
        'RPGゲーミフィケーションという独自のコンセプト（レベル、XP、クラス、クエスト）',
        'フルスタックTypeScript構成による開発・メンテナンス効率',
        '日本市場向け設計（日本語UI、国内HR慣行への対応）',
        'AIチャット統合による日次振り返り・自己成長支援',
        'アバターカスタマイズによる高いエンゲージメント',
        'ピクセルアートUIによる独特のブランドアイデンティティ',
        '低コストなクラウドネイティブアーキテクチャ（React + Express + PostgreSQL）',
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    add_styled_heading(doc, '弱み (Weaknesses)', level=2)
    weaknesses = [
        'モバイルアプリ未対応（Webのみ）',
        '外部システム連携（CRM、Slack、Teams等）が未実装',
        'SSO/SAML対応なし（エンタープライズ要件不足）',
        '高度な分析・レポーティング機能が限定的',
        'マーケットでのブランド認知度が低い',
        'スケーラビリティの検証が不十分',
    ]
    for w in weaknesses:
        doc.add_paragraph(w, style='List Bullet')

    add_styled_heading(doc, '機会 (Opportunities)', level=2)
    opportunities = [
        'HR×ゲーミフィケーション市場の急成長（CAGR 27.9%）',
        '日本国内でのRPG型HRツールの空白市場',
        'リモートワーク普及によるデジタル人材育成需要の増大',
        '経済産業省のゲーミフィケーション人材育成推進',
        'アジア太平洋地域での高成長（CAGR 28.6%）',
        'AI・生成AIとの統合による差別化余地',
    ]
    for o in opportunities:
        doc.add_paragraph(o, style='List Bullet')

    add_styled_heading(doc, '脅威 (Threats)', level=2)
    threats = [
        '大手HRプラットフォーム（HiBob、Lattice等）のゲーミフィケーション機能追加',
        'Centrical等のエンタープライズ向け競合の日本進出',
        '景気後退時のHR Tech投資縮小リスク',
        'ゲーミフィケーションへの一時的ブーム終了リスク',
        '国内大手ゲーム会社（セガ等）のBtoB参入',
    ]
    for t in threats:
        doc.add_paragraph(t, style='List Bullet')

    # ─── 6. 差別化戦略 ───
    add_styled_heading(doc, '6. 差別化戦略と推奨事項')

    doc.add_paragraph(
        'Quest HRの最大の差別化要因は、本格的なRPGメカニクス（キャラクタークラス、'
        'レベルアップ、XPシステム、スキルレーダー、アバターカスタマイズ）を'
        'HR/人材育成に統合している点である。この特徴は現在の競合には見られない独自性であり、'
        '以下の戦略で強化すべきである。'
    )

    recommendations = [
        ('RPGメカニクスの深化', 'パーティ（チーム）クエスト、ギルド（部門）対抗戦、'
         'レアアイテム報酬、シーズンイベント等、RPG要素をさらに拡充し、'
         '継続的なエンゲージメントを促進する。'),
        ('モバイル対応', 'PWA（Progressive Web App）またはReact Nativeでの'
         'モバイルアプリ開発を優先課題とする。従業員の日常的なアクセスを促進。'),
        ('外部連携の拡充', 'Slack、Microsoft Teams、Google Workspace等との連携を実装し、'
         'エンタープライズ導入の障壁を下げる。'),
        ('日本市場特化', '国内HR慣行（目標管理制度、人事考課制度）との連携機能を強化し、'
         '「日本の人事部門のためのゲーミフィケーション」というポジションを確立する。'),
        ('AI機能の強化', 'スキル推薦、キャリアパス提案、パフォーマンス予測等の'
         'AI機能を拡充し、データドリブンな人材育成を実現する。'),
    ]

    for title, desc in recommendations:
        doc.add_paragraph(title, style='List Bullet')
        doc.add_paragraph(desc)

    # ─── 7. 結論 ───
    add_styled_heading(doc, '7. 結論')
    doc.add_paragraph(
        'ゲーミフィケーション×HR市場は急成長しており、Quest HRはRPGメカニクスという'
        '独自の切り口で差別化可能なポジションにある。特に日本国内では、本格的なRPG型'
        'HRプラットフォームは存在せず、ブルーオーシャンの機会がある。'
        'モバイル対応、外部連携、AI強化を優先的に進めることで、'
        '市場での競争優位を確立できると考えられる。'
    )

    filepath = os.path.join(OUTPUT_DIR, '競合調査レポート.docx')
    doc.save(filepath)
    print(f'Created: {filepath}')


def create_revenue_forecast():
    """収益予測レポート"""
    doc = Document()

    title = doc.add_heading('Quest HR 収益予測レポート', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)

    doc.add_paragraph('作成日: 2026年3月1日')
    doc.add_paragraph('対象期間: 2026年4月 ～ 2031年3月（5ヶ年計画）')
    doc.add_paragraph('')

    # ─── 1. エグゼクティブサマリー ───
    add_styled_heading(doc, '1. エグゼクティブサマリー')
    doc.add_paragraph(
        'Quest HRは、RPGゲーミフィケーションを活用した人材育成SaaSプラットフォームとして、'
        'サブスクリプションモデル（月額/年額課金）を主たる収益源とする。'
        '本レポートでは、市場環境、料金体系、顧客獲得見込みに基づく5ヶ年収益予測を示す。'
    )

    # ─── 2. ビジネスモデルと料金体系 ───
    add_styled_heading(doc, '2. ビジネスモデルと料金体系')

    add_styled_heading(doc, '2.1 収益モデル', level=2)
    models = [
        ('SaaSサブスクリプション（主要収益）', '月額/年額のユーザー課金（従量制）'),
        ('導入・カスタマイズ費', '初期設定、カスタマイズ、データ移行等の一時収入'),
        ('プレミアム機能アドオン', 'AI分析、高度レポーティング、API連携等の追加機能'),
        ('研修・コンサルティング', 'ゲーミフィケーション導入コンサルティング、運用研修'),
    ]
    for name, desc in models:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    add_styled_heading(doc, '2.2 料金プラン（想定）', level=2)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    headers = ['プラン', '対象', '月額/ユーザー', '年額/ユーザー', '主な機能']
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], '2D1B69')
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    plans = [
        ['スターター', '～50名', '¥500', '¥5,000', '基本クエスト、レベル/XP、スキル管理'],
        ['スタンダード', '～200名', '¥800', '¥8,000', 'スターター + AIチャット、高度分析、ファイル提出'],
        ['プレミアム', '～1,000名', '¥1,200', '¥12,000', 'スタンダード + SSO、API連携、カスタムクエスト'],
        ['エンタープライズ', '1,000名～', '要相談', '要相談', '全機能 + 専任サポート、SLA、カスタム開発'],
    ]
    for plan in plans:
        row = table.add_row()
        for i, val in enumerate(plan):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ─── 3. 市場規模とTAM/SAM/SOM ───
    add_styled_heading(doc, '3. 市場規模とTAM/SAM/SOM')

    doc.add_paragraph(
        'グローバルのHR向けゲーミフィケーションソフトウェア市場は2025年時点で約12億ドル（約1,800億円）、'
        '日本国内のゲーミフィケーション市場全体は463億円（2024年）と推定されている。'
    )

    table2 = doc.add_table(rows=1, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'

    hdr2 = table2.rows[0].cells
    headers2 = ['区分', '定義', '推定規模']
    for i, h in enumerate(headers2):
        hdr2[i].text = h
        set_cell_shading(hdr2[i], '2D1B69')
        for p in hdr2[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    tam_data = [
        ['TAM（総市場規模）', '国内ゲーミフィケーション×HR市場全体', '約150億円（2026年推定）'],
        ['SAM（対象市場規模）', 'RPG/ゲーム型HR SaaS市場（中堅～大企業）', '約30億円'],
        ['SOM（獲得可能市場）', '初期ターゲット（IT/ゲーム/先進企業）', '約5億円（5年目目標）'],
    ]
    for d in tam_data:
        row = table2.add_row()
        for i, val in enumerate(d):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ─── 4. 顧客獲得計画 ───
    add_styled_heading(doc, '4. 顧客獲得計画')

    table3 = doc.add_table(rows=1, cols=6)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Table Grid'

    hdr3 = table3.rows[0].cells
    headers3 = ['年度', '新規顧客数', '累計顧客数', '平均従業員数/社', '月間チャーン率', 'ARPU（月額）']
    for i, h in enumerate(headers3):
        hdr3[i].text = h
        set_cell_shading(hdr3[i], '2D1B69')
        for p in hdr3[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)

    acquisition = [
        ['1年目 (2026)', '15社', '15社', '80名', '3.0%', '¥640/人'],
        ['2年目 (2027)', '35社', '47社', '120名', '2.5%', '¥720/人'],
        ['3年目 (2028)', '60社', '100社', '150名', '2.0%', '¥800/人'],
        ['4年目 (2029)', '80社', '168社', '180名', '1.8%', '¥900/人'],
        ['5年目 (2030)', '100社', '250社', '200名', '1.5%', '¥1,000/人'],
    ]
    for a in acquisition:
        row = table3.add_row()
        for i, val in enumerate(a):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ─── 5. 5ヶ年収益予測 ───
    add_styled_heading(doc, '5. 5ヶ年収益予測')

    add_styled_heading(doc, '5.1 売上高予測', level=2)

    table4 = doc.add_table(rows=1, cols=7)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    table4.style = 'Table Grid'

    hdr4 = table4.rows[0].cells
    headers4 = ['項目', '1年目', '2年目', '3年目', '4年目', '5年目']
    for i, h in enumerate(headers4):
        hdr4[i].text = h
        set_cell_shading(hdr4[i], '2D1B69')
        for p in hdr4[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)

    revenue_data = [
        ['SaaS月額収入', '¥6.9M', '¥24.4M', '¥57.6M', '¥108.9M', '¥180.0M'],
        ['SaaS年額収入', '¥69.1M', '¥243.6M', '¥576.0M', '¥1,088.6M', '¥1,800.0M'],
        ['導入費（一時）', '¥15.0M', '¥35.0M', '¥60.0M', '¥80.0M', '¥100.0M'],
        ['アドオン収入', '¥0M', '¥12.0M', '¥40.0M', '¥80.0M', '¥150.0M'],
        ['コンサル収入', '¥5.0M', '¥15.0M', '¥30.0M', '¥50.0M', '¥80.0M'],
        ['合計売上高', '¥89.1M', '¥305.6M', '¥706.0M', '¥1,298.6M', '¥2,130.0M'],
    ]
    for d in revenue_data:
        row = table4.add_row()
        for i, val in enumerate(d):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if d[0] == '合計売上高':
                        run.font.bold = True

    doc.add_paragraph('')
    doc.add_paragraph('※ M = 百万円。SaaS年額収入 = SaaS月額収入 × 12ヶ月で概算。')

    add_styled_heading(doc, '5.2 コスト構造と損益予測', level=2)

    table5 = doc.add_table(rows=1, cols=7)
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    table5.style = 'Table Grid'

    hdr5 = table5.rows[0].cells
    for i, h in enumerate(headers4):
        hdr5[i].text = h
        set_cell_shading(hdr5[i], '2D1B69')
        for p in hdr5[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)

    cost_data = [
        ['人件費', '¥36.0M', '¥72.0M', '¥120.0M', '¥180.0M', '¥250.0M'],
        ['インフラ費', '¥6.0M', '¥12.0M', '¥24.0M', '¥40.0M', '¥60.0M'],
        ['マーケティング費', '¥15.0M', '¥40.0M', '¥80.0M', '¥120.0M', '¥150.0M'],
        ['その他経費', '¥12.0M', '¥20.0M', '¥30.0M', '¥40.0M', '¥50.0M'],
        ['合計コスト', '¥69.0M', '¥144.0M', '¥254.0M', '¥380.0M', '¥510.0M'],
        ['営業利益', '¥20.1M', '¥161.6M', '¥452.0M', '¥918.6M', '¥1,620.0M'],
        ['営業利益率', '22.6%', '52.9%', '64.0%', '70.7%', '76.1%'],
    ]
    for d in cost_data:
        row = table5.add_row()
        for i, val in enumerate(d):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if d[0] in ('合計コスト', '営業利益', '営業利益率'):
                        run.font.bold = True

    doc.add_paragraph('')

    # ─── 6. KPI ───
    add_styled_heading(doc, '6. 主要KPI')

    kpis = [
        ('MRR（月次経常収益）', '1年目末: ¥6.9M → 5年目末: ¥180M'),
        ('ARR（年間経常収益）', '1年目末: ¥82.8M → 5年目末: ¥2,160M'),
        ('月間チャーン率', '3.0%から1.5%へ段階的改善'),
        ('CAC（顧客獲得コスト）', '1年目: ¥1.0M/社 → 5年目: ¥1.5M/社'),
        ('LTV（顧客生涯価値）', '1年目: ¥4.6M → 5年目: ¥16.0M'),
        ('LTV/CAC比率', '4.6倍～10.7倍（健全水準3倍以上を維持）'),
        ('NPS（推奨度スコア）', '目標: 40以上（業界平均25前後）'),
    ]
    for name, val in kpis:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(val)

    # ─── 7. リスクと感度分析 ───
    add_styled_heading(doc, '7. リスクと感度分析')

    add_styled_heading(doc, '7.1 主要リスク', level=2)
    risks = [
        ('顧客獲得の遅延', '営業チーム構築・マーケティング施策の成果が想定を下回るリスク。'
         '売上が基本シナリオの60%に留まる可能性あり。'),
        ('チャーン率の悪化', '月間チャーン率が3%を超えると、成長が大幅に鈍化する。'
         'プロダクト品質とカスタマーサクセスの投資が重要。'),
        ('価格競争', '大手プラットフォームの参入により、価格圧力が高まるリスク。'
         '差別化機能（RPGメカニクス）の維持が鍵。'),
        ('開発コストの超過', 'モバイル対応、AI強化等の開発投資が計画を上回るリスク。'),
    ]
    for name, desc in risks:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    add_styled_heading(doc, '7.2 シナリオ分析', level=2)

    table6 = doc.add_table(rows=1, cols=4)
    table6.alignment = WD_TABLE_ALIGNMENT.CENTER
    table6.style = 'Table Grid'

    hdr6 = table6.rows[0].cells
    headers6 = ['シナリオ', '5年目売上高', '5年目営業利益', '前提']
    for i, h in enumerate(headers6):
        hdr6[i].text = h
        set_cell_shading(hdr6[i], '2D1B69')
        for p in hdr6[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    scenarios = [
        ['楽観シナリオ', '¥3,200M', '¥2,500M', '顧客獲得1.5倍、ARPU20%増'],
        ['基本シナリオ', '¥2,130M', '¥1,620M', '本レポートの想定通り'],
        ['悲観シナリオ', '¥850M', '¥300M', '顧客獲得60%、チャーン率1.5倍'],
    ]
    for s in scenarios:
        row = table6.add_row()
        for i, val in enumerate(s):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    # ─── 8. 結論 ───
    add_styled_heading(doc, '8. 結論')
    doc.add_paragraph(
        'Quest HRは、急成長するゲーミフィケーション×HR市場において、'
        'SaaSサブスクリプションモデルにより5年で年間売上21億円を目指す。'
        'RPGメカニクスという独自の差別化要因により、適切なプライシングと'
        '顧客獲得戦略を実行すれば、高い営業利益率（76%）を達成可能である。'
        '初期の顧客獲得とチャーン率の管理が収益性を左右する最重要因子であり、'
        'プロダクト品質とカスタマーサクセスへの投資を惜しまないことが成功の鍵となる。'
    )

    filepath = os.path.join(OUTPUT_DIR, '収益予測レポート.docx')
    doc.save(filepath)
    print(f'Created: {filepath}')


def create_operations_report():
    """運用要件レポート"""
    doc = Document()

    title = doc.add_heading('Quest HR 運用要件レポート', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x2D, 0x1B, 0x69)

    doc.add_paragraph('作成日: 2026年3月1日')
    doc.add_paragraph('～ Quest HRを継続的に運用・成長させるために必要なこと ～')
    doc.add_paragraph('')

    # ─── 1. エグゼクティブサマリー ───
    add_styled_heading(doc, '1. エグゼクティブサマリー')
    doc.add_paragraph(
        '本レポートは、Quest HR（RPG型ゲーミフィケーション人材育成プラットフォーム）を'
        '今後継続的に運用・成長させていくために必要な事項を、技術、組織、ビジネスの'
        '観点から包括的にまとめたものである。'
    )

    # ─── 2. 技術インフラと運用 ───
    add_styled_heading(doc, '2. 技術インフラと運用')

    add_styled_heading(doc, '2.1 現在の技術スタック', level=2)
    doc.add_paragraph('Quest HRは以下の技術スタックで構築されている。')

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['レイヤー', '技術', '備考']
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], '2D1B69')
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    stack = [
        ['フロントエンド', 'React 18 + TypeScript + Vite 7 + TailwindCSS', 'Shadcn UI、Framer Motion'],
        ['バックエンド', 'Express.js 5 + TypeScript', 'Passport.js認証、Multerファイルアップロード'],
        ['データベース', 'PostgreSQL + Drizzle ORM', 'セッションストアもPostgreSQL'],
        ['ビルドツール', 'Vite + ESBuild', 'Hot Module Replacement対応'],
        ['ホスティング', 'Replit（現在）', '本番環境では移行が必要'],
    ]
    for s in stack:
        row = table.add_row()
        for i, val in enumerate(s):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph('')

    add_styled_heading(doc, '2.2 本番環境の構築（最優先）', level=2)
    doc.add_paragraph(
        '現在Replit上で動作しているが、本番運用には以下の環境移行が必要である。'
    )

    infra_items = [
        ('クラウドプラットフォーム選定', 'AWS、GCP、またはAzureへの移行。'
         '推奨: AWS（ECS/Fargate + RDS PostgreSQL + CloudFront）'
         'またはVercel（フロントエンド） + Railway/Render（バックエンド + DB）。'),
        ('コンテナ化', 'Dockerfileの作成。フロントエンド・バックエンド・DBの分離。'
         'docker-composeによるローカル開発環境の構築。'),
        ('CI/CDパイプライン', 'GitHub Actions等による自動ビルド・テスト・デプロイ。'
         'ステージング環境と本番環境の分離。'),
        ('ドメインとSSL', '独自ドメインの取得とSSL証明書の設定（Let\'s Encrypt or ACM）。'),
        ('CDN設定', '静的アセットのCDN配信（CloudFront、Cloudflare等）。'),
    ]
    for name, desc in infra_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    add_styled_heading(doc, '2.3 データベース運用', level=2)
    db_items = [
        'バックアップ: 日次自動バックアップ + ポイントインタイムリカバリ（PITR）の設定',
        'マイグレーション: Drizzle Kitを使った安全なスキーマ変更プロセスの確立',
        'モニタリング: クエリパフォーマンス監視、スロークエリの検出と最適化',
        'スケーリング: リードレプリカの導入計画（ユーザー数増加時）',
        'データ保持ポリシー: チャットログ、クエスト完了データの保持期間と削除ルール',
    ]
    for item in db_items:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '2.4 セキュリティ対策', level=2)
    security_items = [
        'HTTPS強制と適切なセキュリティヘッダー（HSTS、CSP、X-Frame-Options等）',
        'パスワードハッシュ: bcryptは実装済み。パスワードポリシーの強化（最低文字数、複雑性要件）',
        'セッション管理: セッションタイムアウト、同時ログイン制限の検討',
        'CSRF対策: トークンベースのCSRF防御の実装',
        'レート制限: ブルートフォース攻撃防止のためのログイン試行制限',
        'ファイルアップロード: ウイルススキャン、ファイルサイズ・タイプ検証の強化',
        'SQLインジェクション: Drizzle ORMでパラメータ化クエリは実装済み。入力バリデーション強化',
        'XSS対策: Reactの自動エスケープに加え、サニタイゼーション強化',
        '脆弱性スキャン: Dependabot/Snyk等による依存パッケージの自動脆弱性検出',
        'SOC2/ISMS準拠: エンタープライズ顧客向けのセキュリティ認証取得計画',
    ]
    for item in security_items:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '2.5 監視・アラート', level=2)
    monitoring = [
        'アプリケーション監視: Datadog、New Relic、またはGrafana + Prometheusの導入',
        'エラー追跡: Sentry等による実行時エラーのリアルタイム追跡',
        'アップタイム監視: 外部からの死活監視（UptimeRobot等）、SLA 99.9%目標',
        'ログ管理: 構造化ログの導入、ELKスタックまたはCloudWatch Logsでの集約',
        'アラート: PagerDuty/Opsgenie等によるインシデント対応体制の構築',
    ]
    for item in monitoring:
        doc.add_paragraph(item, style='List Bullet')

    # ─── 3. 開発・品質管理 ───
    add_styled_heading(doc, '3. 開発・品質管理')

    add_styled_heading(doc, '3.1 テスト戦略', level=2)
    doc.add_paragraph('現在テストが不足しているため、以下のテスト体制を構築する必要がある。')
    tests = [
        ('ユニットテスト', 'Jest + React Testing Libraryでコンポーネント・関数のテスト。'
         'カバレッジ目標: 80%以上。'),
        ('APIテスト', 'Supertest等によるAPIエンドポイントの自動テスト。'
         '認証、権限チェック、バリデーションを網羅。'),
        ('E2Eテスト', 'Playwright/Cypressによるユーザーフロー全体のテスト。'
         'ログイン→クエスト完了→レベルアップ等の主要フローをカバー。'),
        ('パフォーマンステスト', 'k6/Artillery等による負荷テスト。'
         '同時ユーザー数の上限と応答時間を計測。'),
    ]
    for name, desc in tests:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    add_styled_heading(doc, '3.2 開発プロセス', level=2)
    dev_process = [
        'Git Flow: feature/develop/staging/mainブランチ戦略の確立',
        'コードレビュー: PR必須、最低1名のレビュー承認',
        'リリースサイクル: 2週間スプリント、月次リリース',
        'ドキュメント: API仕様書（OpenAPI/Swagger）、アーキテクチャドキュメント',
        'コーディング規約: ESLint + Prettier設定の統一、TypeScript strictモード',
    ]
    for item in dev_process:
        doc.add_paragraph(item, style='List Bullet')

    # ─── 4. 機能ロードマップ ───
    add_styled_heading(doc, '4. 機能開発ロードマップ')

    add_styled_heading(doc, 'Phase 1: 基盤強化（0～6ヶ月）', level=2)
    phase1 = [
        '本番環境構築（クラウド移行、CI/CD）',
        'テスト基盤の構築（ユニットテスト、E2Eテスト）',
        'PWA対応（モバイルでのWeb体験向上）',
        'パフォーマンス最適化（ページ読み込み速度、API応答時間）',
        'マルチテナント対応（企業別データ分離）',
        'セキュリティ強化（CSRF、レート制限、脆弱性スキャン）',
    ]
    for item in phase1:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, 'Phase 2: 機能拡充（6～12ヶ月）', level=2)
    phase2 = [
        'SSO/SAML対応（エンタープライズ認証）',
        'Slack/Microsoft Teams連携（通知、クエスト完了報告）',
        'パーティ（チーム）クエスト機能',
        '高度な分析ダッシュボード（管理者向け）',
        'カスタムクエストビルダー（管理者がクエストを柔軟に設計）',
        'CSV/Excelインポート/エクスポート機能',
    ]
    for item in phase2:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, 'Phase 3: AI・高度機能（12～24ヶ月）', level=2)
    phase3 = [
        'AIスキル推薦エンジン（従業員の成長に最適なクエストを自動推薦）',
        'AIキャリアパスアドバイザー（長期的なスキル開発計画の提案）',
        'パフォーマンス予測AI（離職リスク、スキルギャップの早期検知）',
        'ギルド（部門）対抗戦・シーズンイベント機能',
        'React Nativeモバイルアプリ',
        'API公開（サードパーティ連携用RESTful API）',
        '多言語対応（英語、中国語等）',
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')

    # ─── 5. 組織体制 ───
    add_styled_heading(doc, '5. 組織体制')

    add_styled_heading(doc, '5.1 必要な人材と役割', level=2)

    table2 = doc.add_table(rows=1, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    headers2 = ['役割', '人数（初期）', '人数（2年目～）', '主な責務', '採用優先度']
    for i, h in enumerate(headers2):
        hdr2[i].text = h
        set_cell_shading(hdr2[i], '2D1B69')
        for p in hdr2[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)

    roles = [
        ['プロダクトマネージャー', '1名', '1～2名', '要件定義、ロードマップ管理、ユーザーリサーチ', '最高'],
        ['フルスタックエンジニア', '2名', '4～6名', '機能開発、バグ修正、コードレビュー', '最高'],
        ['デザイナー（UI/UX）', '1名', '1～2名', 'UI設計、UXリサーチ、デザインシステム管理', '高'],
        ['インフラエンジニア/SRE', '0.5名(兼務)', '1名', 'クラウド構築、監視、セキュリティ、CI/CD', '高'],
        ['QAエンジニア', '0名(兼務)', '1名', 'テスト戦略、自動テスト、品質管理', '中'],
        ['カスタマーサクセス', '1名', '2～3名', '顧客オンボーディング、サポート、フィードバック収集', '高'],
        ['営業/マーケティング', '1名', '2～4名', 'リード獲得、商談、コンテンツマーケティング', '高'],
    ]
    for r in roles:
        row = table2.add_row()
        for i, val in enumerate(r):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph('')

    add_styled_heading(doc, '5.2 初期チーム体制（最小構成: 6名）', level=2)
    doc.add_paragraph('プロダクトマネージャー × 1、フルスタックエンジニア × 2、'
                      'デザイナー × 1、カスタマーサクセス × 1、営業 × 1')

    # ─── 6. カスタマーサクセス ───
    add_styled_heading(doc, '6. カスタマーサクセスと運用サポート')

    add_styled_heading(doc, '6.1 オンボーディングプロセス', level=2)
    onboarding = [
        'キックオフミーティング（ヒアリング、目標設定）',
        '環境セットアップ（テナント作成、初期データ投入、SSO設定）',
        '管理者トレーニング（クエスト設計、従業員管理、分析機能の使い方）',
        '従業員向け説明会支援（アバター作成、初回クエスト体験）',
        '導入後1ヶ月フォロー（利用状況確認、改善提案）',
    ]
    for item in onboarding:
        doc.add_paragraph(item, style='List Bullet')

    add_styled_heading(doc, '6.2 継続的サポート体制', level=2)
    support = [
        'ヘルプセンター/FAQ: よくある質問と操作ガイドのオンラインドキュメント',
        'チャットサポート: 営業時間内のリアルタイムサポート（Intercom等）',
        'メールサポート: 24時間受付、営業日内8時間以内回答',
        'QBR（四半期ビジネスレビュー）: 利用状況分析、改善提案、ロードマップ共有',
        'ユーザーコミュニティ: Slack/Discord等でのユーザー同士の交流の場',
    ]
    for item in support:
        doc.add_paragraph(item, style='List Bullet')

    # ─── 7. 法務・コンプライアンス ───
    add_styled_heading(doc, '7. 法務・コンプライアンス')
    legal = [
        ('個人情報保護法対応', '従業員の個人情報（氏名、スキル、評価データ）を扱うため、'
         '個人情報保護法に準拠したデータ取り扱いポリシーの策定が必須。'
         'プライバシーポリシーの作成と同意取得フローの実装。'),
        ('利用規約', 'SaaS利用規約の策定。データの帰属、サービスレベル、免責事項の明確化。'),
        ('セキュリティ認証', 'SOC2 Type II、ISMS（ISO 27001）の取得検討。'
         'エンタープライズ顧客の信頼獲得に必要。'),
        ('データ所在地', '日本国内のデータセンターでのホスティング。'
         '海外クラウドを使う場合はリージョン選定に注意。'),
        ('労働関連法規', 'ゲーミフィケーションが過度な競争やハラスメントに繋がらないよう、'
         'ガイドラインを提供。匿名化オプションの検討。'),
    ]
    for name, desc in legal:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    # ─── 8. マーケティング・成長戦略 ───
    add_styled_heading(doc, '8. マーケティング・成長戦略')

    strategies = [
        ('コンテンツマーケティング', 'ブログ、ホワイトペーパー、事例紹介を通じた'
         'ゲーミフィケーション×人材育成の啓蒙活動。SEO対策。'),
        ('無料トライアル', '14日間の無料トライアル提供。フリーミアムモデルの検討'
         '（5名まで無料等）。'),
        ('パートナーシップ', 'HR系SaaS（SmartHR、カオナビ等）とのAPI連携・共同マーケティング。'
         '人材育成コンサルティング企業との協業。'),
        ('展示会・カンファレンス', 'HR Tech、Japan IT Week等への出展。'
         'デモブースでのRPG体験型プレゼンテーション。'),
        ('プロダクトレッドグロース', 'プロダクト内での口コミ促進（チームクエスト完了時の共有機能等）。'
         'NPS調査に基づく改善サイクル。'),
        ('リファラルプログラム', '既存顧客からの紹介による新規獲得。紹介者に割引や特典を提供。'),
    ]
    for name, desc in strategies:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(name + ': ')
        run.bold = True
        p.add_run(desc)

    # ─── 9. 優先度マトリクス ───
    add_styled_heading(doc, '9. 優先度マトリクス')

    table3 = doc.add_table(rows=1, cols=4)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    table3.style = 'Table Grid'
    hdr3 = table3.rows[0].cells
    headers3 = ['優先度', '項目', '期限目標', '担当']
    for i, h in enumerate(headers3):
        hdr3[i].text = h
        set_cell_shading(hdr3[i], '2D1B69')
        for p in hdr3[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)

    priorities = [
        ['P0（即時）', '本番環境構築（クラウド移行）', '1ヶ月以内', 'エンジニア/SRE'],
        ['P0（即時）', 'セキュリティ対策強化', '1ヶ月以内', 'エンジニア'],
        ['P0（即時）', '利用規約/プライバシーポリシー策定', '1ヶ月以内', '法務/PM'],
        ['P1（高）', 'テスト基盤構築', '3ヶ月以内', 'エンジニア'],
        ['P1（高）', 'CI/CDパイプライン構築', '3ヶ月以内', 'エンジニア/SRE'],
        ['P1（高）', 'PWAモバイル対応', '3ヶ月以内', 'エンジニア/デザイナー'],
        ['P1（高）', 'カスタマーサクセス体制構築', '3ヶ月以内', 'CS担当'],
        ['P2（中）', 'マルチテナント対応', '6ヶ月以内', 'エンジニア'],
        ['P2（中）', 'SSO/SAML対応', '6ヶ月以内', 'エンジニア'],
        ['P2（中）', '外部サービス連携（Slack/Teams）', '6ヶ月以内', 'エンジニア'],
        ['P2（中）', 'セキュリティ認証（SOC2/ISMS）着手', '6ヶ月以内', '全社'],
        ['P3（低）', 'React Nativeモバイルアプリ', '12～18ヶ月', 'エンジニア'],
        ['P3（低）', 'AI高度機能（予測分析等）', '12～24ヶ月', 'エンジニア/DS'],
        ['P3（低）', '多言語対応', '18～24ヶ月', 'エンジニア/PM'],
    ]
    for pr in priorities:
        row = table3.add_row()
        for i, val in enumerate(pr):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph('')

    # ─── 10. 結論 ───
    add_styled_heading(doc, '10. 結論')
    doc.add_paragraph(
        'Quest HRを成功するSaaSプロダクトとして運用していくには、技術基盤の強化、'
        'セキュリティ対策、テスト体制の構築を最優先で進める必要がある。'
        '同時に、カスタマーサクセス体制の確立と適切なチーム編成により、'
        '顧客満足度の向上とチャーン率の低減を実現する。'
    )
    doc.add_paragraph(
        '短期的には本番環境の安定稼働とセキュリティ確保、'
        '中期的にはモバイル対応と外部連携による利便性向上、'
        '長期的にはAI機能とグローバル展開による差別化を段階的に進めることで、'
        '持続的な成長を実現できる。'
    )

    filepath = os.path.join(OUTPUT_DIR, '運用要件レポート.docx')
    doc.save(filepath)
    print(f'Created: {filepath}')


if __name__ == '__main__':
    print('Creating Quest HR research reports...')
    create_competitive_analysis()
    create_revenue_forecast()
    create_operations_report()
    print('All reports created successfully!')
